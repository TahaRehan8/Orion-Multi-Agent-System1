"""
Word Document Export for Orion
Converts markdown-formatted agent responses to .docx files
"""

import os
import re
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(markdown_text: str, title: str = "Orion Response") -> str:
    """
    Convert markdown text to a Word document and return the temp file path.
    The file is saved to a temp directory for browser download.
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add timestamp
    timestamp = doc.add_paragraph()
    run = timestamp.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # spacer
    
    # Process markdown lines
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 50)
            i += 1
            continue
        
        # Headers
        if line.startswith('### '):
            doc.add_heading(clean_markdown(line[4:]), level=3)
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(clean_markdown(line[3:]), level=2)
            i += 1
            continue
        elif line.startswith('# '):
            doc.add_heading(clean_markdown(line[2:]), level=1)
            i += 1
            continue
        
        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            para = doc.add_paragraph(style='List Bullet')
            add_formatted_text(para, text)
            i += 1
            continue
        
        # Numbered lists
        numbered_match = re.match(r'^\s*(\d+)\.\s+(.*)', line)
        if numbered_match:
            text = numbered_match.group(2)
            para = doc.add_paragraph(style='List Number')
            add_formatted_text(para, text)
            i += 1
            continue
        
        # Table detection
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            # Parse table
            headers = [h.strip() for h in line.split('|') if h.strip()]
            i += 2  # Skip header and separator
            
            rows = []
            while i < len(lines) and '|' in lines[i]:
                row = [c.strip() for c in lines[i].split('|') if c.strip()]
                rows.append(row)
                i += 1
            
            # Create table
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Headers
                for j, header in enumerate(headers):
                    table.rows[0].cells[j].text = clean_markdown(header)
                
                # Data rows
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        if c_idx < len(headers):
                            table.rows[r_idx + 1].cells[c_idx].text = clean_markdown(cell)
                
                doc.add_paragraph()  # spacer after table
            continue
        
        # Regular paragraph
        para = doc.add_paragraph()
        add_formatted_text(para, line)
        i += 1
    
    # Save to temp file
    temp_dir = tempfile.gettempdir()
    filename = f"orion_response_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)
    
    return filepath


def clean_markdown(text: str) -> str:
    """Remove markdown formatting characters"""
    # Remove bold
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove italic
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove inline code
    text = re.sub(r'`(.*?)`', r'\1', text)
    # Remove links [text](url) -> text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    # Remove emoji-like status icons
    text = text.replace('✅', '[PASS]').replace('❌', '[FAIL]')
    text = text.replace('⏳', '').replace('🔄', '').replace('•', '')
    return text.strip()


def add_formatted_text(paragraph, text: str):
    """Add text with bold/italic formatting preserved to a paragraph"""
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif '`' in part:
            # Handle inline code
            code_parts = re.split(r'(`.*?`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = paragraph.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(80, 80, 80)
                else:
                    if cp:
                        paragraph.add_run(cp)
        else:
            if part:
                paragraph.add_run(part)
